from oscopilot.tool_repository.basic_tools.base_action import BaseAction
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt, Inches
import os
import re
from tqdm import tqdm

class create_cheat_sheet(BaseAction):
    def __init__(self):
        self._description = "Convert multiple PDF files into a condensed Word document cheat sheet."

    def clean_text(self, text):
        """Clean unnecessary characters and content from text"""
        # Remove special characters and extra whitespace
        text = re.sub(r'[\n\r\t]', ' ', text)
        text = re.sub(r'\s+', ' ', text)
        
        # Remove common unnecessary content
        patterns_to_remove = [
            r'\d+/\d+',  # Page numbers
            r'©.*?All Rights Reserved',  # Copyright information
            r'www\..*?\.com',  # Website addresses
            r'http\S+',  # URLs
            r'Figure\s*\d+',  # Figure labels
            r'Table\s*\d+',  # Table labels
        ]
        
        for pattern in patterns_to_remove:
            text = re.sub(pattern, '', text)
        
        return text.strip()

    def __call__(self, working_directory=None, pdf_folder='Slides', output_docx='output.docx', *args, **kwargs):
        """
        Convert multiple PDF files into a condensed Word document cheat sheet.

        Args:
            pdf_folder (str): Path to the folder containing PDF files. Default is 'Slides'.
            output_docx (str): Path for the output Word document. Default is 'output.docx'.

        Returns:
            None
        """
        # Create new Word document
        doc = Document()
        
        # Set margins (0.3 inches)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.3)
            section.bottom_margin = Inches(0.3)
            section.left_margin = Inches(0.3)
            section.right_margin = Inches(0.3)
        
        # Initial font size and line spacing
        font_size = 6.5
        line_spacing = 0.8
        
        while True:
            doc = Document()  # Recreate document
            
            # Reset margins
            for section in doc.sections:
                section.top_margin = Inches(0.3)
                section.bottom_margin = Inches(0.3)
                section.left_margin = Inches(0.3)
                section.right_margin = Inches(0.3)
            
            # Process all PDF files
            for filename in tqdm(os.listdir(pdf_folder)):
                if filename.lower().endswith('.pdf'):
                    pdf_path = os.path.join(pdf_folder, filename)
                    pdf_reader = PdfReader(pdf_path)
                    
                    text = ""
                    for page in tqdm(pdf_reader.pages, desc="Reading pages", leave=False):
                        text += page.extract_text()
                    
                    # Only clean the text, not the summary
                    text = self.clean_text(text)
                    
                    paragraph = doc.add_paragraph(text)
                    paragraph.paragraph_format.line_spacing = line_spacing
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)
            
            # Temporary save the document for checking the page count
            temp_path = 'temp.docx'
            doc.save(temp_path)
            temp_doc = Document(temp_path)
            
            # Check the page count
            if len(temp_doc.sections) <= 2:
                doc.save(output_docx)
                os.remove(temp_path)
                break
            
            # If more than 2 pages, decrease the font size or line spacing
            if font_size > 4:  # Minimum font size limit
                font_size -= 0.5
            elif line_spacing > 0.5:  # Minimum line spacing limit
                line_spacing -= 0.1
            else:
                print("Warning: Unable to compress content to less than 2 pages!")
                doc.save(output_docx)
                os.remove(temp_path)
                break

# Example usage
# create_cheat_sheet(pdf_folder='Slides', output_docx='output.docx')
